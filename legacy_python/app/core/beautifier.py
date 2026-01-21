import re
import os
import logging
from docx import Document
from docx.shared import Pt, Inches

log = logging.getLogger(__name__)

def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

def add_bold_run(paragraph, text_before, text_bold, text_after=""):
    paragraph.add_run(text_before)
    paragraph.add_run(text_bold).bold = True
    if text_after: paragraph.add_run(text_after)

def run_beautifier(input_path, output_path):
    log.info(f"Beautifying {input_path} -> {output_path}")
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        doc = Document()
        set_style(doc)
        bold_regex = re.compile(r'\*\*(.*?)\*\*')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line: continue
            
            if line.startswith('# Class Report:'):
                doc.add_heading(line.replace('# Class Report:', '').strip(), 0)
                if i > 5: doc.add_page_break()
                continue
            if line.startswith('## '):
                doc.add_heading(line.replace('## ', '').strip(), 1)
                continue
            if line.startswith('### '):
                doc.add_heading(line.replace('### ', '').strip(), 2)
                continue
                
            p = None
            if line.startswith('* **'): 
                line = line.lstrip('* ')
                p = doc.add_paragraph(style='List Bullet')
            elif line.startswith('* '):
                line = line.lstrip('* ')
                p = doc.add_paragraph(style='List Bullet')
            elif line.startswith('    * '):
                line = line.lstrip('    * ')
                p = doc.add_paragraph(style='List Bullet 2')
                p.paragraph_format.left_indent = Inches(0.5)
            else:
                p = doc.add_paragraph()

            match = bold_regex.search(line)
            if match:
                s, e = match.span()
                add_bold_run(p, line[:s], match.group(1), line[e:])
            else:
                p.add_run(line)

        doc.save(output_path)
        return True, output_path
    
    except Exception as e:
        return False, str(e)
